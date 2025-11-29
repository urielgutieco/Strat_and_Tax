import os
import io
import zipfile
import random
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Inches

# ----------------- Drive con SERVICE ACCOUNT (compatible con Vercel) -----------------
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
# -------------------------------------------------------------------------------------

app = Flask(__name__)
CORS(app)

# Carpetas raíz
TEMPLATE_FOLDER = 'template_word'
GENERATED_DOCS = 'template_gendocs'
GENERATED_ZIPS = 'template_genzips'

os.makedirs(TEMPLATE_FOLDER, exist_ok=True)
os.makedirs(GENERATED_DOCS, exist_ok=True)
os.makedirs(GENERATED_ZIPS, exist_ok=True)

# Mapeo servicio → carpeta
SERVICIO_TO_DIR = {
    "Servicios de construccion de unidades unifamiliares": "construccion_unifamiliar",
    "Servicios de reparacion o ampliacion o remodelacion de viviendas unifamiliares": "reparacion_remodelacion_unifamiliar",
    "Servicio de remodelacion general de viviendas unifamiliares": "remodelacion_general",
    "Servicios de reparacion de casas moviles en el sitio": "reparacion_casas_moviles",
    "Servicios de construccion y reparacion de patios y terrazas": "patios_terrazas",
    "Servico de reparacion por daños ocasionados por fuego de viviendas unifamiliares": "reparacion_por_fuego",
    "Servicio de construccion de casas unifamiliares nuevas": "construccion_unifamiliar_nueva",
    "Servicio de instalacion de casas unifamiliares prefabricadas": "instalacion_prefabricadas",
    "Servicio de construccion de casas en la ciudad o casas jardin unifamiliares nuevas": "construccion_casas_ciudad_jardin",
    "Dasarrollo urbano": "desarrollo_urbano",
    "Servicio de planificacion de la ordenacion urbana": "planificacion_ordenacion_urbana",
    "Servicio de administracion de tierras urbanas": "administracion_tierras_urbanas",
    "Servicio de programacion de inversiones urbanas": "programacion_inversiones_urbanas",
    "Servicio de reestructuracion de barrios marginales": "reestructuracion_barrios_marginales",
    "Servicios de alumbrado urbano": "alumbrado_urbano",
    "Servicios de control o regulacion del desarrollo urbano": "control_desarrollo_urbano",
    "Servicios de estandares o regulacion de edificios urbanos": "estandares_regulacion_edificios",
    "Servicios comunitarios urbanos": "comunitarios_urbanos",
    "Servicios de administracion o gestion de proyectos o programas urbanos": "gestion_proyectos_programas_urbanos",
    "Ingenieria civil": "ingenieria_civil",
    "Ingenieria de carreteras": "ingenieria_carreteras",
    "Ingenieria deinfraestructura de instalaciones o fabricas": "infraestructura_instalaciones_fabricas",
    "Servicios de mantenimiento e instalacion de equipo pesado": "mantenimiento_instalacion_equipo_pesado",
    "Servicio de mantenimiento y reparacion de equipo pesado": "mantenimiento_reparacion_equipo_pesado",
}

# Plantillas
TEMPLATE_FILES = [
    'plantilla_solicitud.docx',
    '2.docx',
    '3.docx',
    '4.docx',
    '1.docx',
]

def replace_text_in_document(document, replacements):
    for paragraph in document.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))

def generate_single_document(template_filename, template_root, replacements, user_image_path=None, data=None):
    template_path = os.path.join(template_root, template_filename)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Plantilla '{template_filename}' no encontrada en '{template_root}'.")

    document = Document(template_path)
    replace_text_in_document(document, replacements)

    # Imagen
    if user_image_path and os.path.exists(user_image_path):
        try:
            document.add_paragraph()
            document.add_paragraph(
                data.get('nombre_completo_de_la_persona_que_firma_la_solicitud', 'N/A')
                if data else 'N/A'
            )
            document.add_picture(user_image_path, width=Inches(2.5))
        except:
            document.add_paragraph("⚠ No se pudo insertar la imagen del usuario.")
    else:
        document.add_paragraph("⚠ Imagen de firma no encontrada en el servidor.")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# ------------------------ NUEVA FUNCIÓN PARA GOOGLE DRIVE (SERVICE ACCOUNT) ------------------------

def authenticate_and_upload_to_drive(file_name, zip_buffer):
    """
    Autenticación con Google Drive usando Service Account (compatible con Vercel)
    """

    try:
        service_account_info = os.getenv("-----BEGIN PRIVATE KEY-----\nMIIEvQIBADANBgkqhkiG9w0BAQEFAASCBKcwggSjAgEAAoIBAQDAmGjF74jEfxg1\ngUF910V/BOg/rdKrLD8mfdcZ6Lr4cAu8B1zOTkzNSw1lmg1bA5W1ugNF7/YxGf/Z\nMMxkaSsme1WAUVgNe/dx7qPK8Q1wrCGiSFDnaEbd/uQH3/r6rfvlM0oUte9fs4ts\n6I9FsdQjwhmwMJzJgOld+Jbr/C+oGWJbv3y8J9iarkD5a1ZMKLx5ZiP1wJvd+wpC\n/DMnc3PLrA9tFc7pogPkwx6x6p34zG+n++1RKmmf225t0XHVgI7UqRr9tlqXYjC7\nowq5VfqWO8a5qY3YiaaB+gd/wQ+0ZNRna+8M89gJYZSBBNYPGl/E4Pj+Nud5zUSU\n/WS3id/lAgMBAAECggEAExiJeuavJtx7KALy3WlUyK1R4c54TuSKNmTPyQ/cSfnC\nEBiiyQnWb5x7bZGCtX17gJM689pBDMlJJt3BxkrvLleYOKrYEi4ycKd6sgqIjJ1k\nkFSfQPKD5Er7jRRWLDjH+wfE1pzbaR/COUAtxHlHUWxTY7bONOSFPv4A6NEOYJpH\npEJrw14vm/8lpJZes9Bgp0noXt0msNZlJk5n2FBIEqn837bO7tO1x1NtqMoMW9ng\n8f/zW0U3PlQOOMSRuW/pBYBnXKXkXFaqx0M+KH/h7gvQttFnhnoroy+glW3colSE\nL4cnKLeDP0wpC1bKsEsN+UrBsA9sSk1n9jKaduNmGQKBgQDobZXfrrNmX3/Ega06\nGIdmiSVbbIn587QoiMV3dalRuzTUKyKIfgpmjTGfhULdVhM80dYGKTRJrby3Fryy\nEJnp6KYxDhELkCeZwgdp9z3btqMamkTPhw12VOqm/h5D///SlO8tfCo30wZ9onVg\nbDVa0Rh35ZTWIEOaW07e0rTxWwKBgQDUIKqtIn6mnyQFV6xUsHmvJVZ3CHvUqJlh\nB69wVu+Lc+aVuWnpna331b95LXRpps+zv1dpHgLKULMkM0cUKB3IFkRvm3HlRz9a\nJjRGE7j//TuWWcZa4ral/2OJmE+WWpp80SyEbU3XutS7pXELTRk/t/NAI2JPXvyJ\nEaHj1fH3vwKBgQCzR6v+IGwiv/D2qyvqDveJ4KmfDmaTFSbWyUC/d1OGRodmTGtT\nqxzso2YubIT54yXtmNGkaO37EvboprIgC2wdH6XpWrdNGwFguWgslVfyfLrsjyga\nh0qcBr4E1yiTPQif7t9aT1blPnHYQJKXMIArL/PWr5CRZcufaWS5vP9y1wKBgHB6\nOcW/6qPy8iibCleFk6AZDjXjm7VxCJL4fj/0+ONau0Ncoxoqb5pgDjz0qytyNeO7\na1/jd9tK2xIw1lmLw+7aT8NWCxhlpOPqKgGWZ5vk7HmWdKSwXeS6/E5zIpA4zn2s\ndAxr4MCo4i/1U3GImgC9SxARMwme9gy/+rVSNhEFAoGAKR0JkhUX4hR60Yps9wvk\n16aFFR9acmEDDaMFgOGHT2E4sT6hlvsTvATcB+ik7+v5aZSjL+A8foQXhXWwjUlt\nwzEwe1CvH54075n/P3yySjiJW760SpNuQlXwdvHjoWg1ntKTB97EwyjlEge48W/s\neCBHYNWtJT4OYgPgGSDGR4Q=\n-----END PRIVATE KEY-----\n")

        if not service_account_info:
            return {"success": False, "message": "Falta variable de entorno GOOGLE_SERVICE_ACCOUNT_JSON"}

        creds = Credentials.from_service_account_info(
            eval(service_account_info),
            scopes=["https://www.googleapis.com/auth/drive.file"]
        )

        service = build('drive', 'v3', credentials=creds)

        # Metadata
        file_metadata = {
            'name': file_name
        }

        zip_buffer.seek(0)
        media = MediaIoBaseUpload(zip_buffer, mimetype='application/zip', resumable=True)

        uploaded_file = service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id'
        ).execute()

        return {
            "success": True,
            "message": f"Archivo subido correctamente. ID: {uploaded_file.get('id')}"
        }

    except Exception as e:
        return {"success": False, "message": f"Error al subir a Drive: {str(e)}"}

# -----------------------------------------------------------------------------------------------


@app.route('/generate-word', methods=['POST'])
def generate_word():
    try:
        data = request.form.to_dict()

        uploaded_image = request.files.get("imagen_usuario")
        user_image_path = None

        if uploaded_image:
            user_image_path = os.path.join(TEMPLATE_FOLDER, "imagen_custom.png")
            uploaded_image.save(user_image_path)

        if not data:
            return jsonify({"error": "No data received"}), 400

        for filename in os.listdir(GENERATED_DOCS):
            file_path = os.path.join(GENERATED_DOCS, filename)
            try:
                if os.path.isfile(file_path):
                    os.remove(file_path)
            except Exception as e:
                print(f"Error al eliminar {file_path}: {e}")

        servicio = data.get('servicio')
        if not servicio:
            return jsonify({"error": "Debes seleccionar un servicio."}), 400

        carpeta_servicio = SERVICIO_TO_DIR.get(servicio)
        if not carpeta_servicio:
            return jsonify({"error": f"No existe carpeta mapeada para el servicio: {servicio}"}), 404

        template_root = os.path.join(TEMPLATE_FOLDER, carpeta_servicio)
        if not os.path.isdir(template_root):
            return jsonify({"error": f"La carpeta de plantillas no existe: {template_root}"}), 404

        numero_de_contrato_unico = ''.join([str(random.randint(0, 9)) for _ in range(18)])
        descripcion_servicio = servicio

        replacements = {
            '${descripcion_del_servicio}': descripcion_servicio,
            '${razon_social}': data.get('razon_social', 'N/A'),
            '${r_f_c}': data.get('r_f_c', 'N/A'),
            '${domicilio_del_cliente}': data.get('domicilio_del_cliente', 'N/A'),
            '${telefono_del__cliente}': data.get('telefono_del__cliente', 'N/A'),
            '${correo_del_cliente}': data.get('correo_del_cliente', 'N/A'),
            '${fecha_de_inicio_del_servicio}': data.get('fecha_de_inicio_del_servicio', 'N/A'),
            '${fecha_de_conclusion_del_servicio}': data.get('fecha_de_conclusion_del_servicio', 'N/A'),
            '${monto_de_la_operacion_Sin_IVA}': data.get('monto_de_la_operacion_Sin_IVA', 'N/A'),
            '${forma_de_pago}': data.get('forma_de_pago', 'N/A'),
            '${cantidad}': data.get('cantidad', 'N/A'),
            '${unidad}': data.get('unidad', 'N/A'),
            '${numero_de_contrato}': numero_de_contrato_unico,
            '${fecha_de_operación}': data.get('fecha_de_operación', 'N/A'),
            '${nombre_completo_de_la_persona_que_firma_la_solicitud}': data.get('nombre_completo_de_la_persona_que_firma_la_solicitud', 'N/A'),
            '${cargo_de_la_persona_que_firma_la_solicitud}': data.get('cargo_de_la_persona_que_firma_la_solicitud', 'N/A'),
            '${factura_relacionada_con_la_operación}': data.get('factura_relacionada_con_la_operación', 'N/A'),
            '${informe_si_cuenta_con_fotografias_videos_o_informacion_adicion}': data.get('informe_si_cuenta_con_fotografias_videos_o_informacion_adicion', 'N/A'),
            '${comentarios}': data.get('comentarios', 'N/A')
        }

        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for template in TEMPLATE_FILES:
                try:
                    doc_buffer = generate_single_document(template, template_root, replacements, user_image_path, data)
                    base = os.path.splitext(template)[0]

                    rfc = data.get('r_f_c', 'N/A')
                    output_filename = f"{base}_{descripcion_servicio}_{base}_{numero_de_contrato_unico}_{rfc}.docx"

                    output_path = os.path.join(GENERATED_DOCS, output_filename)
                    if os.path.exists(output_path):
                        os.remove(output_path)

                    with open(output_path, "wb") as f:
                        f.write(doc_buffer.getvalue())

                    zip_file.writestr(output_filename, doc_buffer.getvalue())

                except Exception as e:
                    print(f"Error generando documento {template}: {e}")
                    continue

        zip_buffer.seek(0)

        final_zip_name = f"{descripcion_servicio}_{numero_de_contrato_unico}_{data.get('r_f_c', 'N/A')}.zip"

        zip_server_path = os.path.join(GENERATED_ZIPS, final_zip_name)
        with open(zip_server_path, "wb") as zip_disk:
            zip_disk.write(zip_buffer.getvalue())

        # ---------- SUBIR A DRIVE ----------
        zip_buffer.seek(0)
        upload_result = authenticate_and_upload_to_drive(final_zip_name, zip_buffer)

        print(f"GOOGLE DRIVE: {upload_result['message']}")

        zip_buffer.seek(0)

        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=final_zip_name
        )

    except Exception as e:
        return jsonify({"error": f"Error interno: {str(e)}"}), 500


if __name__ == '__main__':
    app.run(debug=True, port=5001)
